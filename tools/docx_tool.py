"""DOCX tool — generate Word documents."""
import os, uuid

NAME        = "docx_tool"
DESCRIPTION = "Generate a Microsoft Word (.docx) document from text content"
CATEGORY    = "builtin"
ICON        = "📝"
INPUTS = [
    {"name": "content",  "label": "Document Content", "type": "textarea",
     "placeholder": "Enter the document content...", "required": True},
    {"name": "title",    "label": "Document Title",   "type": "text",
     "placeholder": "My Document"},
    {"name": "filename", "label": "Output Filename",  "type": "text",
     "placeholder": "output.docx"},
]

DOCS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "generated_docs")


def run(content: str, title: str = "Document", filename: str = "") -> dict:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        os.makedirs(DOCS_DIR, exist_ok=True)
        if not filename:
            filename = f"doc_{uuid.uuid4().hex[:8]}.docx"
        if not filename.endswith(".docx"):
            filename += ".docx"
        path = os.path.join(DOCS_DIR, filename)

        doc = Document()

        # Title
        if title:
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Body — parse markdown-style headings
        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif not stripped:
                doc.add_paragraph("")
            else:
                # Handle **bold**
                p = doc.add_paragraph()
                import re
                parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

        doc.save(path)
        return {"message": f"✅ Word document generated!\n[📥 Download {filename}](/docs/{filename})",
                "file": f"/docs/{filename}"}

    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        return {"error": f"DOCX generation failed: {e}"}
